# -*- coding: utf-8 -*-
"""Rebuild the "פירוש צדקה אל-חכים" commentary (tzdaka_sections / tzdaka_verse_links)
from the full set of manuscript docx files, extending coverage beyond Gen 1–6.

Sources, in priority order (later file wins on a duplicate ref):
  1. tzdaka_bereshit_alef_bet_39.docx        — chapters 1–10 (supersedes the old
                                               perek_alef_10 import; its 1–6 == it)
  2. tzdaka_bereshit_yod-bet_yod-chet.docx   — chapters 12–18 (… up to 18:20)
  3. tzdaka_bereshit_18-20.docx              — 18:20 → 20 (headingless: ref-lines
                                               appear as ordinary paragraphs)
(yod-bet_yod-zayin_1.docx is a 12–17 subset of #2 and is intentionally skipped.)

Section format (same as the original importer):
  Heading 2 / ref-line:  'א:ב — „incipit” · topic'  (ranges use '–'; numerals may
                         carry gershayim/geresh, e.g. 'י״ב:א–ג', 'י״ח:כ״ג').
  Body = the paragraphs until the next section.
Headings that aren't verse refs (פרק…, הערת מתרגם) close the current section;
once a 'נספח'/appendix heading is reached the file is no longer scanned.

Each section links to the verses in its ref (book בראשית). The incipit is validated
against the linked verse text. Rebuilds ONLY the two tzdaka tables; full DB backup
before --apply.

Usage:  py -3 scripts/import_tzdaka_all.py            # dry run
        py -3 scripts/import_tzdaka_all.py --apply
"""
import sqlite3, sys, io, os, re, shutil, datetime

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
APPLY = '--apply' in sys.argv
DB = 'data/torah.db'
BOOK = 'בראשית'
TZDIR = 'תרגום צדקה אלכים'        # the source docx were moved into this subfolder
FILES = [os.path.join(TZDIR, n) for n in (
    'tzdaka_bereshit_alef_bet_39.docx',
    'tzdaka_bereshit_yod-bet_yod-chet.docx',
    'tzdaka_bereshit_18-20.docx',
    # ch 43-46 use the ref-line format ("מג:א — „incipit” …") handled by parse_file
    'sadaqah_gen43_full.docx', 'sadaqah_gen44_full.docx',
    'sadaqah_gen45_full.docx', 'sadaqah_gen46_full.docx')]
# manuscripts in the "(verse-number) <verse text>  + commentary paragraphs" format
# (Gen 21–30). The "SP_aligned" ch21 files are skipped — their commentary is an
# OCR-failure placeholder, superseded by the full ch21 file. gen25b_26 spans two
# chapters (25 & 26).
MARKER_FILES = [os.path.join(TZDIR, n) for n in (
    'sadaqah_gen21_full_1.docx', 'sadaqah_gen22_full.docx',
    'sadaqah_gen23_full.docx', 'sadaqah_gen24_full.docx',
    'sadaqah_gen25_full_1.docx',                 # ch25:1-11
    'sadaqah_gen25b_26_full.docx',               # ch25:29-34 + ch26
    'sadaqah_gen27_full.docx', 'sadaqah_gen28_full.docx',
    'sadaqah_gen29_full.docx', 'sadaqah_gen30_full.docx')]
# Gen 47-50: Heading-1 'פרק <heb>' sets the chapter, Heading-2 is a vocalised verse
# incipit ending in a comma-bracketed ref '[מז, א]', then the commentary paragraphs.
PART5_FILES = [os.path.join(TZDIR, 'part5_continued_ch47-50.docx')]
# Gen 31-40: same layout as part5 but the Heading-2 verse incipit ends in a COLON-
# bracketed ref '[לא: א]' (not a comma), and some headings carry the start of the
# commentary after the ref — that trailing text is pushed into the body.
PART4_FILES = [os.path.join(TZDIR, 'part4_ch31-40.docx')]

GEM = {'א':1,'ב':2,'ג':3,'ד':4,'ה':5,'ו':6,'ז':7,'ח':8,'ט':9,'י':10,'כ':20,
       'ך':20,'ל':30,'מ':40,'ם':40,'נ':50,'ן':50,'ס':60,'ע':70,'פ':80,'ף':80,
       'צ':90,'ץ':90,'ק':100,'ר':200,'ש':300,'ת':400}
NIK = re.compile('[֑-ׇ]')
NUM = r'[א-ת]+(?:[״׳][א-ת]*)*'                       # numeral, gershayim/geresh ok
REF_RE = re.compile(r'^(%s)\s*:\s*(%s(?:[–\-]%s)?)' % (NUM, NUM, NUM))
# a standalone ref-line in a headingless file: ref + separator (· or —) + incipit
LINE_RE = re.compile(r'^(%s)\s*:\s*(%s(?:[–\-]%s)?)\s*[·—]\s*[„"]' % (NUM, NUM, NUM))
# an authoritative bracketed ref at the line end, e.g. '… [מה:כד]'; the gen43-46
# files carry it and it overrides a typo in the leading ref.
BRACK_RE = re.compile(r'\[(%s)\s*:\s*(%s(?:[–\-]%s)?)\]' % (NUM, NUM, NUM))
# the Gen 47-50 (part5_continued) headings end with a COMMA-separated bracketed ref,
# e.g. '… [מז, א]' (chapter, verse) — not a colon.
BRACK_COMMA = re.compile(r'\[(%s)\s*[,،]\s*(%s(?:[–\-]%s)?)\]' % (NUM, NUM, NUM))


def gem(s):
    return sum(GEM.get(c, 0) for c in s)


def bare(w):
    return NIK.sub('', (w or '')).replace('־', ' ')


def parse_ref(h):
    """('chap', [verse_ints], incipit, topic, ref) from a heading/ref-line, or None."""
    m = REF_RE.match(h.strip())
    if not m:
        return None
    rch, rvs = m.group(1), m.group(2)
    mb = BRACK_RE.search(h)                 # a bracketed ref overrides the leading one
    if mb:
        rch, rvs = mb.group(1), mb.group(2)
    ch = gem(rch)
    vs = re.split(r'[–\-]', rvs)
    v1 = gem(vs[0]); v2 = gem(vs[1]) if len(vs) > 1 and gem(vs[1]) else v1
    rest = h[m.end():]
    inc = ''
    mi = re.search(r'[„"]([^”"]+)[”"]', rest)
    if mi:
        inc = mi.group(1).strip()
    topic = rest.split('·', 1)[1].strip() if '·' in rest else ''
    # strip a trailing bracketed ref from the topic so it isn't shown as the title
    topic = re.sub(r'\s*\[%s\s*:\s*%s(?:[–\-]%s)?\]\s*$' % (NUM, NUM, NUM), '', topic).strip()
    ref = '%s:%s' % (rch, rvs)
    return ch, list(range(v1, v2 + 1)), inc, topic, ref


# "(<gematria>) <verse text>" — the verse number may be a range, e.g. (ה-ו)
VERSE_RE = re.compile(r'^\(\s*([א-ת]{1,3})(?:\s*[-–]\s*([א-ת]{1,3}))?\s*\)\s*(.*)')
# a chapter heading line: 'בראשית · פרק כ״ו' or 'פרק כ״ו' (anchored, so prose
# mentions of "פרק" don't trigger). Files may span several chapters (e.g. 25 & 26).
HEAD_RE = re.compile(r'^(?:בראשית\s*[·]\s*)?פרק\s+(%s)' % NUM)


def parse_marker_file(path, vidx):
    """Parse a manuscript where each verse is introduced by '(<gematria>) <verse
    text>' and the following paragraphs are its commentary (Gen 21–30). The chapter
    is tracked from 'פרק <num>' headings and may change within a file."""
    import docx
    doc = docx.Document(path)
    sections = []; cur = None; chap = None; chap_he = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        plain = bare(t)
        mh = HEAD_RE.match(plain)
        if mh and 'פרקים' not in plain:                    # chapter heading (transition)
            c = gem(mh.group(1))
            if 1 <= c <= 50:
                chap, chap_he = c, mh.group(1)
            continue                                       # not a verse, not body
        mv = VERSE_RE.match(plain)
        if mv and chap:
            if cur:
                sections.append(cur)
            v1 = gem(mv.group(1)); v2 = gem(mv.group(2)) if mv.group(2) else v1
            vns = list(range(v1, v2 + 1))
            vids = [vidx[(chap, vn)] for vn in vns if (chap, vn) in vidx]
            ref_v = mv.group(1) + ('-' + mv.group(2) if mv.group(2) else '')
            cur = {'ch': chap, 'ref': '%s:%s' % (chap_he, ref_v), 'title': '',
                   'incipit': mv.group(3).strip(), 'vids': vids,
                   'missing': [vn for vn in vns if (chap, vn) not in vidx], 'body': []}
        elif cur is not None:
            cur['body'].append(t)
    if cur:
        sections.append(cur)
    out = []
    for s in sections:                                     # drop OCR-failure placeholders
        body = ' '.join(s['body'])
        if 'לא ניתן לשחזור' in body or not body.strip():
            continue
        out.append(s)
    return out


def parse_part5_file(path, vidx):
    """Gen 47-50: 'פרק <heb>' (Heading 1) sets the chapter; each Heading-2 is a verse
    incipit ending with a comma-bracketed ref '[מז, א]', and the paragraphs after it
    are its commentary. The bracket ref (chapter+verse) is authoritative."""
    import docx
    doc = docx.Document(path)
    sections = []; cur = None; chap = None; chap_he = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        sty = (p.style.name if p.style else '') or ''
        plain = bare(t)
        mh = HEAD_RE.match(plain)
        if sty.startswith('Heading 1') and mh and 'פרקים' not in plain:
            c = gem(mh.group(1))
            if 1 <= c <= 50:
                chap, chap_he = c, mh.group(1)
            continue
        mb = BRACK_COMMA.search(t)
        if sty.startswith('Heading 2') and mb:
            if cur:
                sections.append(cur)
            rch = gem(mb.group(1)); rvs = mb.group(2)
            vs = re.split(r'[–\-]', rvs)
            v1 = gem(vs[0]); v2 = gem(vs[1]) if len(vs) > 1 and gem(vs[1]) else v1
            vns = list(range(v1, v2 + 1))
            vids = [vidx[(rch, vn)] for vn in vns if (rch, vn) in vidx]
            incipit = BRACK_COMMA.sub('', t).strip(' ·—-')
            cur = {'ch': rch, 'ref': '%s:%s' % (mb.group(1), rvs), 'title': '',
                   'incipit': incipit, 'vids': vids,
                   'missing': [vn for vn in vns if (rch, vn) not in vidx], 'body': []}
        elif cur is not None and not sty.startswith('Heading 1'):
            cur['body'].append(t)
    if cur:
        sections.append(cur)
    return sections


def parse_part4_file(path, vidx):
    """Gen 31-40: 'פרק <heb>' (Heading 1) sets the chapter; each Heading-2 is a verse
    incipit ending with a COLON-bracketed ref '[לא: א]', and the paragraphs after it
    are its commentary. Any commentary that trails the ref on the heading line itself
    is moved into the body so it isn't lost."""
    import docx
    doc = docx.Document(path)
    sections = []; cur = None; chap = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        sty = (p.style.name if p.style else '') or ''
        plain = bare(t)
        mh = HEAD_RE.match(plain)
        if sty.startswith('Heading 1') and mh and 'פרקים' not in plain:
            c = gem(mh.group(1))
            if 1 <= c <= 50:
                chap = c
            continue
        mb = BRACK_RE.search(t)
        if sty.startswith('Heading 2') and mb:
            if cur:
                sections.append(cur)
            rch = gem(mb.group(1)); rvs = mb.group(2)
            vs = re.split(r'[–\-]', rvs)
            v1 = gem(vs[0]); v2 = gem(vs[1]) if len(vs) > 1 and gem(vs[1]) else v1
            vns = list(range(v1, v2 + 1))
            vids = [vidx[(rch, vn)] for vn in vns if (rch, vn) in vidx]
            incipit = t[:mb.start()].strip(' ·—-')          # text before the ref
            cur = {'ch': rch, 'ref': '%s:%s' % (mb.group(1), rvs), 'title': '',
                   'incipit': incipit, 'vids': vids,
                   'missing': [vn for vn in vns if (rch, vn) not in vidx], 'body': []}
            post = t[mb.end():].strip(' ·—-')               # commentary trailing the ref
            if post:
                cur['body'].append(post)
        elif cur is not None and not sty.startswith('Heading 1'):
            cur['body'].append(t)
    if cur:
        sections.append(cur)
    return sections


def parse_file(path, vidx):
    """Yield section dicts from one docx (heading-based or headingless)."""
    import docx
    doc = docx.Document(path)
    sections = []
    cur = None
    for p in doc.paragraphs:
        sty = (p.style.name if p.style else '') or ''
        t = p.text.strip()
        if not t:
            continue
        if t.startswith('נספח') and sty.startswith('Heading'):
            break             # the נספח א׳/ב׳ appendices — stop (NOT 'מסה נספחת')
        is_ref = None
        if sty == 'Heading 2':
            is_ref = parse_ref(t)
        elif LINE_RE.match(t):                       # headingless ref-line
            is_ref = parse_ref(t)
        if sty.startswith('Heading') or is_ref:
            if cur:
                sections.append(cur); cur = None
            if is_ref:
                ch, vns, inc, topic, ref = is_ref
                vids = [vidx[(ch, vn)] for vn in vns if (ch, vn) in vidx]
                cur = {'ch': ch, 'ref': ref, 'title': topic, 'incipit': inc,
                       'vids': vids, 'missing': [vn for vn in vns if (ch, vn) not in vidx],
                       'body': []}
        elif cur is not None:
            cur['body'].append(t)
    if cur:
        sections.append(cur)
    return sections


def main():
    conn = sqlite3.connect(DB, timeout=60); conn.row_factory = sqlite3.Row
    vidx, vtext = {}, {}
    for r in conn.execute(
            """SELECT v.id vid, c.number ch, v.number vn, v.text txt
               FROM verses v JOIN chapters c ON c.id=v.chapter_id
               JOIN books b ON b.id=c.book_id WHERE b.name=?""", (BOOK,)):
        if str(r['vn']).isdigit():
            vidx[(r['ch'], int(r['vn']))] = r['vid']
            vtext[r['vid']] = r['txt'] or ''

    # gather sections. A ref may legitimately repeat WITHIN a file (two distinct
    # comments on the same verse — keep both). Across files the same ref is the same
    # passage in another edition — the LATER file replaces the earlier (e.g. 18:20:
    # file #3's brief note is replaced by file #1's dedicated treatment).
    final = {}          # ref -> list[section]  (from the last file that defines it)
    order = []
    jobs = ([(f, parse_file) for f in FILES] + [(f, parse_marker_file) for f in MARKER_FILES]
            + [(f, parse_part4_file) for f in PART4_FILES]
            + [(f, parse_part5_file) for f in PART5_FILES])
    for f, parser in jobs:
        if not os.path.exists(f):
            print('MISSING file:', f); continue
        secs = parser(f, vidx)
        print('%-46s -> %d sections (chapters %s)'
              % (os.path.basename(f), len(secs), sorted(set(s['ch'] for s in secs))))
        grouped = {}
        for s in secs:
            s['text'] = re.sub(r'\s+', ' ', ' '.join(s['body'])).strip()
            grouped.setdefault(s['ref'], []).append(s)
        for ref, slist in grouped.items():
            if ref not in final:
                order.append(ref)
            final[ref] = slist           # last file wins; within-file keeps all
    sections = [s for ref in order for s in final[ref]]

    # validate incipit against linked verse text
    problems = []
    for s in sections:
        if not s['incipit'] or not s['vids']:
            if s['missing']:
                problems.append((s['ref'], 'missing verses %s' % s['missing']))
            continue
        iw = [w for w in bare(s['incipit']).split() if len(w) >= 3]
        joined = ' '.join(bare(vtext[v]) for v in s['vids'])
        hit = sum(1 for w in iw if w in joined)
        if hit / max(1, len(iw)) < 0.4:
            problems.append((s['ref'], 'incipit "%s" weak match (%d/%d)'
                             % (s['incipit'][:28], hit, len(iw))))

    nlinks = sum(len(s['vids']) for s in sections)
    print('\nTOTAL sections: %d   verse-links: %d   chapters: %s'
          % (len(sections), nlinks, sorted(set(s['ch'] for s in sections))))
    print('validation issues (%d):' % len(problems))
    for ref, msg in problems:
        print('  [%s] %s' % (ref, msg))
    print('\nsample (new chapters):')
    for s in sections:
        if s['ch'] >= 12:
            print('  %s · %s' % (s['ref'], s['title'][:40]))
            print('     %s' % s['text'][:90])
            if s['ch'] >= 13:
                break

    if not APPLY:
        print('\n[dry-run] re-run with --apply to write'); conn.close(); return

    bak = DB + '.bak_tzdaka_all_' + datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
    shutil.copy2(DB, bak); print('backed up ->', os.path.basename(bak))
    cu = conn.cursor()
    cu.execute('DROP TABLE IF EXISTS tzdaka_verse_links')
    cu.execute('DROP TABLE IF EXISTS tzdaka_sections')
    cu.execute('CREATE TABLE tzdaka_sections (id INTEGER PRIMARY KEY, book TEXT, '
               'chap INTEGER, ref TEXT, title TEXT, ord INTEGER, text TEXT)')
    cu.execute('CREATE TABLE tzdaka_verse_links (id INTEGER PRIMARY KEY, '
               'verse_id INTEGER, section_id INTEGER)')
    for i, s in enumerate(sections):
        if not s['vids'] or not s['text']:
            continue
        cu.execute('INSERT INTO tzdaka_sections (book, chap, ref, title, ord, text) '
                   'VALUES (?,?,?,?,?,?)', (BOOK, s['ch'], s['ref'], s['title'], i, s['text']))
        sid = cu.lastrowid
        for vid in s['vids']:
            cu.execute('INSERT INTO tzdaka_verse_links (verse_id, section_id) VALUES (?,?)',
                       (vid, sid))
    cu.execute('CREATE INDEX ix_tzl_verse ON tzdaka_verse_links (verse_id)')
    conn.commit()
    ns = conn.execute('SELECT COUNT(*) FROM tzdaka_sections').fetchone()[0]
    nl = conn.execute('SELECT COUNT(*) FROM tzdaka_verse_links').fetchone()[0]
    nv = conn.execute('SELECT COUNT(DISTINCT verse_id) FROM tzdaka_verse_links').fetchone()[0]
    print('APPLIED: %d sections, %d links across %d verses' % (ns, nl, nv))
    conn.close()


if __name__ == '__main__':
    main()
