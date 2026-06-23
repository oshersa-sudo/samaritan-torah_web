# -*- coding: utf-8 -*-
"""Import the responsa of Jacob ben Aaron the High Priest — "שו"ת של יעקב בן אהרן
הכהן" (from ספר ההארה) — as a new "ממקור שומרון" source. The work is 25 questions;
Part B holds each question's text, and Part C is a biblical index (5 tables, one per
book) mapping verses → question. We link each question to the verses the index ties
to it. Additive tables shyt_sections / shyt_verse_links only.

The same questions also surface under the "מן המסורת השומרונית" (eyalk) source — that
is done in get_eyalk_commentary (UNION), not by duplicating rows here.

Usage:  py -3 scripts/import_shyt.py            # dry run
        py -3 scripts/import_shyt.py --apply
"""
import sqlite3, sys, io, os, re, shutil, datetime

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
APPLY = '--apply' in sys.argv
DOCX = os.path.join('ספר ההארות של יעקב בן אהרן',
                    'ספר_ההארה_ניתוח_תרגום_ואינדקס_מקראי.docx')
DB = 'data/torah.db'
GEM = {'א':1,'ב':2,'ג':3,'ד':4,'ה':5,'ו':6,'ז':7,'ח':8,'ט':9,'י':10,'כ':20,
       'ך':20,'ל':30,'מ':40,'ם':40,'נ':50,'ן':50,'ס':60,'ע':70,'פ':80,'ף':80,
       'צ':90,'ץ':90,'ק':100,'ר':200,'ש':300,'ת':400}
BOOKS = ['בראשית', 'שמות', 'ויקרא', 'במדבר', 'דברים']   # table order in Part C
ISOLATE = re.compile('[⁦-⁩]')                  # bidi isolates around numbers
HEAD_RE = re.compile(r'^פרק\s+([א-ת]+)\s*\(([IVXLC]+)\)\s*·\s*(.+)$')
QCH_RE = re.compile(r'פרק\s+([א-ת]+)')


def gem(s):
    return sum(GEM.get(c, 0) for c in s)


def process_body(body):
    """Tidy a question body and split off its biblical-anchor tail.
    Returns (text, anchors): 'הר גריזים' is written as one word 'הרגריזים', and the
    "עיגון מקראי:" list at the end is separated so it can be shown small, at the foot."""
    body = re.sub(r'הר[\s־]+גריזים', 'הרגריזים', body)
    i = body.find('עיגון מקראי')
    if i >= 0:
        return body[:i].strip(' ·•—-'), body[i:].strip()
    return body.strip(), ''


def parse_questions(doc):
    """{qnum: {'heb','title','body'}} from the Part-B 'פרק <heb> (<roman>) · topic'."""
    qs = {}
    cur = None
    for p in doc.paragraphs:
        sty = (p.style.name if p.style else '') or ''
        t = p.text.strip()
        if not t:
            continue
        if sty == 'Heading 2':
            m = HEAD_RE.match(t)
            if m:
                heb, topic = m.group(1), m.group(3).strip()
                cur = {'heb': heb, 'qnum': gem(heb),
                       'title': 'שאלה %s · %s' % (heb, topic), 'body': []}
                qs[gem(heb)] = cur
            else:
                cur = None              # a Part-A / index heading — stop collecting
        elif cur is not None:
            cur['body'].append(t)
    return qs


def parse_index(doc, vidx):
    """list of (verse_id, qnum) from the 5 per-book index tables."""
    links = []; missing = []
    for ti, table in enumerate(doc.tables[:5]):
        book = BOOKS[ti]
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 2 or 'פסוק' in cells[0]:        # header row
                continue
            ref = ISOLATE.sub('', cells[0]).strip()
            mq = QCH_RE.search(cells[1])
            if not mq or ':' not in ref:
                continue
            qnum = gem(mq.group(1))
            ch, vs = ref.split(':', 1)
            try:
                ch = int(ch)
            except ValueError:
                continue
            parts = re.split(r'[–\-]', vs)
            try:
                v1 = int(parts[0]); v2 = int(parts[1]) if len(parts) > 1 else v1
            except ValueError:
                continue
            for vn in range(v1, v2 + 1):
                vid = vidx.get((book, ch, vn))
                if vid:
                    links.append((vid, qnum))
                else:
                    missing.append('%s %d:%d' % (book, ch, vn))
    return links, missing


def main():
    import docx
    doc = docx.Document(DOCX)
    conn = sqlite3.connect(DB, timeout=60); conn.row_factory = sqlite3.Row
    vidx = {}
    for r in conn.execute("""SELECT v.id vid, b.name bk, c.number ch, v.number vn
                             FROM verses v JOIN chapters c ON c.id=v.chapter_id
                             JOIN books b ON b.id=c.book_id"""):
        if str(r['vn']).isdigit():
            vidx[(r['bk'], r['ch'], int(r['vn']))] = r['vid']

    qs = parse_questions(doc)
    links, missing = parse_index(doc, vidx)
    # keep only links whose question we actually parsed
    links = [(vid, q) for (vid, q) in links if q in qs]
    by_q = {}
    for vid, q in links:
        by_q.setdefault(q, set()).add(vid)

    print('questions parsed:', len(qs), '(qnums %s)' % sorted(qs))
    print('verse-links:', sum(len(v) for v in by_q.values()),
          'across', len(set(v for s in by_q.values() for v in s)), 'verses')
    if missing:
        print('index refs not in DB (%d):' % len(missing), missing[:12])
    print('\nsample:')
    for q in sorted(qs)[:3]:
        body = re.sub(r'\s+', ' ', ' '.join(qs[q]['body']))
        print('  [%d] %s -> %d verses' % (q, qs[q]['title'], len(by_q.get(q, []))))
        print('     %s' % body[:90])

    if not APPLY:
        print('\n[dry-run] re-run with --apply to write'); conn.close(); return

    bak = DB + '.bak_shyt_' + datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
    shutil.copy2(DB, bak); print('backed up ->', os.path.basename(bak))
    cu = conn.cursor()
    cu.execute('DROP TABLE IF EXISTS shyt_verse_links')
    cu.execute('DROP TABLE IF EXISTS shyt_sections')
    cu.execute('CREATE TABLE shyt_sections (id INTEGER PRIMARY KEY, qnum INTEGER, '
               'title TEXT, ord INTEGER, text TEXT, anchors TEXT)')
    cu.execute('CREATE TABLE shyt_verse_links (id INTEGER PRIMARY KEY, '
               'verse_id INTEGER, section_id INTEGER)')
    for q in sorted(qs):
        body = re.sub(r'\s+', ' ', ' '.join(qs[q]['body'])).strip()
        if not body or not by_q.get(q):
            continue
        text, anchors = process_body(body)
        cu.execute('INSERT INTO shyt_sections (qnum, title, ord, text, anchors) VALUES (?,?,?,?,?)',
                   (q, qs[q]['title'], q, text, anchors))
        sid = cu.lastrowid
        for vid in sorted(by_q[q]):
            cu.execute('INSERT INTO shyt_verse_links (verse_id, section_id) VALUES (?,?)',
                       (vid, sid))
    cu.execute('CREATE INDEX ix_shyt_verse ON shyt_verse_links (verse_id)')
    conn.commit()
    ns = conn.execute('SELECT COUNT(*) FROM shyt_sections').fetchone()[0]
    nl = conn.execute('SELECT COUNT(*) FROM shyt_verse_links').fetchone()[0]
    nv = conn.execute('SELECT COUNT(DISTINCT verse_id) FROM shyt_verse_links').fetchone()[0]
    print('APPLIED: %d sections, %d links across %d verses' % (ns, nl, nv))
    conn.close()


if __name__ == '__main__':
    main()
