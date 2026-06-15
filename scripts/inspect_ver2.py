# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn

d = Document(r"data\torah_aziz_ver2.docx")
print("paragraphs:", len(d.paragraphs), "tables:", len(d.tables), "sections:", len(d.sections))
for i, s in enumerate(d.sections):
    cols = s._sectPr.xpath("./w:cols")
    num = cols[0].get(qn('w:num')) if cols else None
    print("  section", i, "cols:", num)

# total text length, count of chapter markers and stops
full = "\n".join(p.text for p in d.paragraphs)
print("total chars:", len(full))
print("'--:' count:", full.count('--:'), " ':--' count:", full.count(':--'))
print("'.' count:", full.count('.'), " ':' count:", full.count(':'))

print("\n=== first 12 non-empty paragraphs ===")
cnt = 0
for p in d.paragraphs:
    t = p.text.strip()
    if t:
        print(f'[{len(t)}c] {t[:150]!r}')
        cnt += 1
        if cnt >= 12:
            break
