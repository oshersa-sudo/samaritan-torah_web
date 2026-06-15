# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

d = Document(r"data\torah_aziz.docx")
print("paragraphs:", len(d.paragraphs))
print("tables:", len(d.tables))
print("sections:", len(d.sections))
for i, s in enumerate(d.sections):
    cols = s._sectPr.xpath("./w:cols")
    num = cols[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num') if cols else None
    print("  section", i, "num cols:", num)

cnt = 0
for p in d.paragraphs:
    t = p.text.strip()
    if t:
        print(repr(t[:140]))
        cnt += 1
        if cnt >= 40:
            break
